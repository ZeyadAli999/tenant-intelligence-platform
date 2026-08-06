"""DOCX paragraphs and tables without macro or embedded-content execution."""

from pathlib import Path

from docx import Document

from services.documents.parsers.base import ParsedDocument, ParsedElement
from services.documents.upload_security import FileValidationError


class DOCXParser:
    def parse(self, path: Path) -> ParsedDocument:
        try:
            document = Document(path)
            elements: list[ParsedElement] = []
            heading: str | None = None
            for paragraph in document.paragraphs:
                text = paragraph.text.strip()
                if not text:
                    continue
                if paragraph.style and paragraph.style.name.startswith("Heading"):
                    heading = text
                elements.append(ParsedElement(text=text, section_title=heading))
            for table_index, table in enumerate(document.tables, 1):
                rows = [
                    " | ".join(cell.text.strip() for cell in row.cells)
                    for row in table.rows
                ]
                text = "\n".join(row for row in rows if row.strip(" |"))
                if text:
                    elements.append(
                        ParsedElement(text=text, section_title=f"Table {table_index}")
                    )
            return ParsedDocument(tuple(elements))
        except Exception as exc:
            raise FileValidationError("PARSING_FAILED") from exc
